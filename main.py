import math
import random
import os
from flask import Flask, flash, request, redirect, url_for, render_template, send_file
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt
from flask_pymongo import PyMongo


UPLOAD_FOLDER = 'upload/'

app = Flask(__name__)


app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MONGO_URI'] = "mongodb://localhost:27017/myDatabase"

mongo = PyMongo(app)

def allowed_doc(filename):
    return '.' in filename and \
        filename.rsplit('.', 1)[1].lower() == 'docx'


def allowed_picture(filename):
    return '.' in filename and \
        filename.rsplit('.', 1)[1].lower() == 'jpg'


@app.route('/')
def home_page():
    return render_template('index.html')


# @app.route('/testbank')
# def bank():
@app.route('/uploaddb', methods=['GET', 'POST'])
def upload_db():
    if request.method == 'POST':
        # check if the post request has the file part
        if 'file' not in request.files:
            return render_template('dbupload.html', msg='No file selected')
        file = request.files['file']
        # if user does not select file, browser also
        # submit an empty part without filename
        if file.filename == '':
            return render_template('dbupload.html', msg='No file selected')

        # if it is a doc
        numClasses= request.form["num_classes"]
        nameTest = request.form["test_title"]
        if file and allowed_doc(file.filename) and numClasses and nameTest:
            filename = secure_filename(file.filename)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))

            for x in range(int(numClasses)):
                doc_to_final(file.filename, nameTest, x)

            return render_template('dbupload.html',
                                   msg='Successfully processed', doc_src='static/' + nameTest + '.docx')
            return redirect(url_for('view_tests()'))

    return render_template("dbupload.html")

@app.route('/testbank')
def view_tests():
    directory = os.fsencode("DB")
    filelist=[]
    for file in os.listdir(directory):
        filename = os.fsdecode(file)
        newFileName="../DB/" + filename
        filelist.append(newFileName)
        print(newFileName)
    unit = filelist[0]
    

    return render_template("testbank.html", unit = unit[10], list = filelist)


    
@app.route('/DB/<address>', methods = ['GET'])
def download_file(address):
    if request.method == 'GET':
        fileAddress = "DB/" + address
        return send_file(fileAddress)
    

@app.route('/upload', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        # check if the post request has the file part
        if 'file' not in request.files:
            return render_template('upload.html', msg='No file selected')
        file = request.files['file']
        # if user does not select file, browser also
        # submit an empty part without filename
        if file.filename == '':
            return render_template('upload.html', msg='No file selected')

        # if it is a doc
        numTest = request.form["num_copies"]
        nameTest = request.form["test_title"]
        if file and allowed_doc(file.filename) and numTest and nameTest:
            filename = secure_filename(file.filename)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))

            doc_to_doc(filename, int(numTest), nameTest)
            return render_template('upload.html',
                                   msg='Successfully processed', doc_src='static/' + nameTest + '.docx')
           # return redirect(url_for('upload_file()'))

    return render_template("upload.html")


if __name__ == '__main__':
    app.run(debug=True)


class MCQuestion:
    def __init__(self, question, choice_list, correct_ans):
        self.question = question
        self.choice_list = choice_list
        self.correct_ans = correct_ans


class FRQuestion:
    def __init__(self, question):
        self.question = question


class Test:
    def __init__(self, mcList, frList):
        self.mcList = mcList
        self.frList = frList


def get_random_number(num_string):
    num = float(num_string)

    lower_bound = 10**(int(math.log10(num)))
    upper_bound = 10**(int(math.log10(num)) + 1)

    random_num = random.uniform(lower_bound, upper_bound)

    random_num_string = str(random_num)
    return random_num_string[0:len(num_string)]


def randomize_frq(frq):
    test_list = frq.question.split()

    for num in range(len(test_list)):
        if '[' in test_list[num]:
            num_string = test_list[num][1:-1]
            random_num = get_random_number(num_string)
            test_list[num] = random_num

    new_frq = frq
    new_frq.question = " ".join(test_list)
    return new_frq


def randomize_mcq(mcq):
    correct_ans_string = mcq.choice_list[mcq.correct_ans]

    test_list = mcq.choice_list
    random.shuffle(test_list)

    new_mcq = mcq
    new_mcq.choice_list = test_list
    new_mcq.correct_ans = test_list.index(correct_ans_string)
    return new_mcq


def randomize_test(test):
    new_test = test

    for num in range(len(new_test.mcList)):
        new_test.mcList[num] = randomize_mcq(new_test.mcList[num])

    for num in range(len(new_test.frList)):
        new_test.frList[num] = randomize_frq(new_test.frList[num])

    new_mc = new_test.mcList
    random.shuffle(new_mc)
    new_test.mcList = new_mc

    new_fr = new_test.frList
    random.shuffle(new_fr)
    new_test.frList = new_fr

    return new_test


def parse_test(test_text):
    questions = test_text.split("\n\n")
    test_return = Test([], [])

    for question in questions:
        q_list = question.split("\n")
        if len(q_list) > 1:
            question_parts = q_list[0].split('. ')
            question_parts.pop(0)
            the_question = ". ".join(question_parts)
            the_choices = []
            correct_choice = 100
            for num in range(1, len(q_list)):
                the_choices.append(q_list[num].split('. ')[1])
                if "*" in q_list[num]:
                    correct_choice = num - 1

            test_return.mcList.append(MCQuestion(
                the_question, the_choices, correct_choice))

        else:
            question_parts = q_list[0].split('. ')
            question_parts.pop(0)
            the_question = ". ".join(question_parts)
            test_return.frList.append(FRQuestion(the_question))

    return test_return


def parse_document(filename):
    print("--------------------------------- " + filename +
          " ------------------------------------------------------")
    document = Document('upload/' + filename)

    test_string = ""
    for p in document.paragraphs:
        test_string += p.text + "\n"

    # print(test_string)
    return parse_test(test_string[:-1])


def doc_to_final(filename, test_name, index):
    letter_list = ['A', 'B', 'C', 'D', 'E', 'F',
                   'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O']
    original_test = parse_document(filename)

    document = Document()

    print("///////////////////////////////////////////////////////")
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri (Body)'
    font.size = Pt(12)

    
    answer_key_list = []
    document.add_heading(test_name + ' Class ' + str(index+1))
    the_test = randomize_test(original_test)
    current_question = 1
    for mcq in the_test.mcList:
        document.add_paragraph(str(current_question) + ". " + mcq.question)
        current_question += 1
        current_choice = 0
        for choice in mcq.choice_list:
            document.add_paragraph(
                '\t' + letter_list[current_choice] + '. ' + choice)
            current_choice += 1

        answer_key_list.append(letter_list[mcq.correct_ans])
        document.add_paragraph()

    for frq in the_test.frList:
        document.add_paragraph(str(current_question) + ". " + frq.question)
        current_question += 1
        document.add_paragraph()

    document.add_page_break()

    document.add_heading(test_name + ' Class ' + str(index+1))

    for num in range(len(answer_key_list)):
        document.add_paragraph(str(num + 1) + ". " + answer_key_list[num])

    document.save('DB/' + test_name + 'class' + str(index+1) + '.docx')

def doc_to_doc(filename, num_copies, test_name):

    letter_list = ['A', 'B', 'C', 'D', 'E', 'F',
                   'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O']
    original_test = parse_document(filename)

    document = Document()

    print("///////////////////////////////////////////////////////")
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri (Body)'
    font.size = Pt(12)

    for num in range(num_copies):
        answer_key_list = []
        document.add_heading(test_name + " Form " + str(num + 1), 0)
        the_test = randomize_test(original_test)
        current_question = 1
        for mcq in the_test.mcList:
            document.add_paragraph(str(current_question) + ". " + mcq.question)
            current_question += 1
            current_choice = 0
            for choice in mcq.choice_list:
                document.add_paragraph(
                    '\t' + letter_list[current_choice] + '. ' + choice)
                current_choice += 1

            answer_key_list.append(letter_list[mcq.correct_ans])
            document.add_paragraph()

        for frq in the_test.frList:
            document.add_paragraph(str(current_question) + ". " + frq.question)
            current_question += 1
            document.add_paragraph()

        document.add_page_break()

        document.add_heading("Answer Key Form " + str(num + 1))

        for num in range(len(answer_key_list)):
            document.add_paragraph(str(num + 1) + ". " + answer_key_list[num])

        document.add_page_break()

    document.save('static/' + test_name + '.docx')
# mongo.save_file(test_name + '.docx',)
